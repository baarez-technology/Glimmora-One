"""Generate the six tech-deliverable Word documents for Glimmora ONE.

Outputs:
  docs/Tech docs/01_USP.docx
  docs/Tech docs/02_Target_Audience.docx
  docs/Tech docs/03_Solution_Design.docx
  docs/Tech docs/04_Project_Documents.docx
  docs/Tech docs/05_Test_Cases.docx
  docs/Tech docs/06_Technical_Documentation.docx

Reflects the state of `main` as of v0.5.0 — slim MVP + redesigned 4-role
hierarchy + subscriptions + notifications + creator-application pipeline
+ in-app Glimmora Guide chatbot + chrome-less floating shell + Postgres.

Run from repo root:
    python scripts/build_tech_docs.py
"""

from __future__ import annotations

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Tech docs"
OUT.mkdir(parents=True, exist_ok=True)

PROJECT = "Glimmora ONE"
VERSION = "0.5.0"
TODAY = datetime.now().strftime("%Y-%m-%d")
OWNER = "Glimmora — sanjay@glimmora.ai"


# ---------------------- helpers ----------------------

def _setup(doc: Document) -> None:
    for s in doc.sections:
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)


def title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x33, 0x2B, 0x2B)
    p.paragraph_format.space_after = Pt(2)


def subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    p.paragraph_format.space_after = Pt(14)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x33, 0x2B, 0x2B)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xAD, 0x6B, 0x2D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x44, 0x3C, 0x3C)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(11)


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)


def cover(doc, doc_title: str, deliverable: str) -> None:
    title(doc, PROJECT)
    subtitle(doc, f"{deliverable}  ·  v{VERSION}  ·  {TODAY}  ·  {OWNER}")
    h1(doc, doc_title)


def table(doc, header: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = val
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)


def page_break(doc):
    doc.add_page_break()


# ============================================================
# 1. USP
# ============================================================

def build_usp():
    doc = Document(); _setup(doc)
    cover(doc, "Unique Selling Proposition", "Deliverable 01 — USP")

    h2(doc, "One-line positioning")
    body(doc, "Glimmora ONE is the calm room of the AI era — a single, attentive space where an "
              "AI companion that listens, a library of wisdom journeys, a quiet reflection "
              "journal, and an in-app help guide hold each other together. Not a tool. A place "
              "to slow down.")

    h2(doc, "Why we exist")
    body(doc, "Most consumer AI is built to maximize engagement: fast replies, infinite scroll, "
              "constant nudges. Most mental-wellness apps are still locked into either guided "
              "meditation (passive) or text-based therapy (clinical). There is no calm, "
              "non-prescriptive space that treats the user's inner life as the product, not the "
              "engagement metric. Glimmora ONE is that space — with a hand-vetted creator marketplace "
              "and a structured admin layer so the calm holds at scale.")

    h2(doc, "Six pillars of the USP")

    h3(doc, "1. An AI that listens first")
    body(doc, "The Companion is trained to be present — short replies, soft language, emotion "
              "classification, and a single 'question to sit with' instead of advice. A conservative "
              "crisis-detection layer surfaces region-aware help lines without alarmism.")

    h3(doc, "2. Wisdom, not content")
    body(doc, "Stories is an OTT-style library of short, beautifully made guided episodes on "
              "stillness, becoming, and emotion. Every episode ends with one quiet question, not a "
              "next-up auto-play. Creators (therapists, monks, teachers, filmmakers) apply via a "
              "public form, are vetted by human moderators, and become creators only after approval.")

    h3(doc, "3. A journal that becomes a digital twin")
    body(doc, "Reflections compound. Glimmora aggregates them into a private map of inner weather "
              "— a 30-day emotional trend chart, dominant moods, streaks, AI-generated 'noticings' "
              "on each entry. Users own it.")

    h3(doc, "4. An in-app guide, not a help-doc trail")
    body(doc, "A floating amber bubble (Glimmora Guide) lives on every page. Left-click opens a "
              "warm, grounded help chat that knows the app inside-out — every flow, every role, "
              "every tier. Right-click + drag repositions it anywhere on the viewport. Users never "
              "have to leave the page to find out how something works.")

    h3(doc, "5. The whole loop in one product")
    body(doc, "Companion, Stories, Journal, Daily Ritual, Digital Twin, Guide — under one calm "
              "shell, one auth, one design language. No tab-switching between an AI app, a "
              "meditation app, a journaling app, and a help portal.")

    h3(doc, "6. Trust at the operational layer")
    body(doc, "Four named roles (Customer, Creator, Moderator, Superadmin) with explicit "
              "permissions. Moderators are human reviewers with their own dashboard. Subscriptions "
              "are admin-granted records with start/end dates. Audit-friendly by design.")

    h2(doc, "Differentiation matrix")
    table(doc,
          header=["Capability", "Headspace / Calm", "BetterHelp", "ChatGPT", "Glimmora ONE"],
          rows=[
              ["AI companion (in-app)", "—", "—", "Yes", "Yes (calm, gated, crisis-aware)"],
              ["Guided video library", "Yes", "—", "—", "Yes (one-question close)"],
              ["Personal reflection journal", "Light", "Inside chat", "—", "Yes (first-class)"],
              ["Digital twin / mood trend", "Light", "—", "—", "Yes (30-day trend + AI noticings)"],
              ["Hand-vetted creator marketplace", "—", "—", "—", "Yes (public apply → moderator review)"],
              ["Crisis safety surface", "Static page", "Therapist", "Generic", "Inline + region-aware"],
              ["In-app help bot grounded in the product", "—", "—", "—", "Yes (Glimmora Guide)"],
              ["Tone", "Polished", "Clinical", "Generic chat", "Calm, soft, present-tense"],
          ])

    h2(doc, "Brand promise")
    body(doc, "We do not chase your attention. We hold space for it.")

    h2(doc, "Tagline")
    body(doc, "“A calm room for your inner life.”")

    out = OUT / "01_USP.docx"; doc.save(str(out)); print(f"wrote {out}")


# ============================================================
# 2. TARGET AUDIENCE
# ============================================================

def build_target_audience():
    doc = Document(); _setup(doc)
    cover(doc, "Target Audience", "Deliverable 02 — Target Audience")

    h2(doc, "Audience summary")
    body(doc, "Glimmora ONE is built for adults — primarily 22–45 — who already journal, meditate, "
              "or work with a therapist, and who are reaching for an integrated inner-work space "
              "that respects their time and intelligence. Geographic priority: India, US, UK, and "
              "the broader English-speaking diaspora; secondary: SEA + EU. Device priority: mobile "
              "web first, desktop web second, native apps later.")

    h2(doc, "Primary personas")

    h3(doc, "1.  The Seeker  (Customer — free / Standard tier)")
    bullet(doc, "Age 22–35, urban, knowledge worker. Has tried Calm/Headspace, journals occasionally, follows a few thoughtful Substacks.")
    bullet(doc, "Pain: feels rushed, restless, slightly disconnected. Doesn't want another optimisation app.")
    bullet(doc, "Success: builds a 1–2× a week habit; uses the Companion as a sounding board; finishes 1–2 series.")
    bullet(doc, "Conversion driver: the free Standard tier is genuinely useful; the daily ritual is the hook.")

    h3(doc, "2.  The Practitioner  (Customer — Premium tier)")
    bullet(doc, "Age 28–45, often in helping professions (therapy, coaching, design, education). Already in long-term therapy or has a meditation practice.")
    bullet(doc, "Pain: scattered tools (Notion + Day One + Calm + ChatGPT). Wants one calm shell.")
    bullet(doc, "Success: long streak; uses the digital twin to notice arcs; references the Companion across weeks.")
    bullet(doc, "Conversion driver: today Premium is admin-granted (no self-serve checkout yet); the Standard→Premium upgrade path is a near-term roadmap item.")

    h3(doc, "3.  The Creator")
    bullet(doc, "Age 30–55. Therapist, monk, teacher, filmmaker, somatic practitioner. Has an audience of 1k–50k.")
    bullet(doc, "Pain: existing platforms (YouTube, Substack, Insight Timer) push virality or reduce them to a thumbnail.")
    bullet(doc, "Success: a meaningful audience that finishes their content.")
    bullet(doc, "How they reach us: click 'Apply as creator' on landing/footer or 'Apply to become a creator' on dashboard → /apply form (full name, email, username, password, pitch, links, attachments). Submission creates the account AND a pending application. Until a moderator decides, they're gated on /under-review.")

    h3(doc, "4.  The Moderator")
    bullet(doc, "Glimmora staff or trusted volunteer. Created only by a Superadmin via the /admin/moderators panel.")
    bullet(doc, "Job: open /moderate/applications, read each pitch + links + attachments by hand, write an optional note, Approve or Reject. Approval flips the applicant's role to Creator and notifies them.")
    bullet(doc, "Success: catches off-brand applicants early; never has to rush.")

    h3(doc, "5.  The Superadmin  (Owner)")
    bullet(doc, "Founder / platform owner. Bootstrapped from env vars on first boot; cannot be created from the UI.")
    bullet(doc, "Sees three role-locked tabs (Customers / Creators / Moderators) with clickable filter tiles. Can edit / delete any non-superadmin account, grant subscriptions, create moderators.")
    bullet(doc, "Success: full visibility into who's on the platform; can rescue, promote, demote, or remove anyone.")

    h2(doc, "Anti-personas (we are not for you, and that is okay)")
    bullet(doc, "Optimisation maximalists who want streak-shaming, leaderboards, and gamification.")
    bullet(doc, "Acute-crisis users who need clinical intervention — Glimmora surfaces helplines and recommends professional care; it is not a substitute.")
    bullet(doc, "Influencers seeking viral reach — there's no like count, no algorithmic feed, no public follower graph.")
    bullet(doc, "Teenagers under 18 — current scope is adults.")

    h2(doc, "Segment sizing (rough)")
    table(doc,
          header=["Segment", "TAM (global)", "SAM (priority geos)", "SOM (Y1)"],
          rows=[
              ["Mindfulness / meditation app users", "~120M", "~40M", "30k"],
              ["Active digital journalers", "~25M", "~9M", "10k"],
              ["AI companion early adopters", "~80M", "~25M", "20k"],
              ["Creators in wisdom / wellbeing", "~250k", "~70k", "120 vetted"],
          ])

    h2(doc, "Jobs to be done (JTBD)")
    bullet(doc, "When my mind is loud and I'm alone, I want to be heard without being judged.")
    bullet(doc, "When I notice a pattern (stuck, anxious, restless) repeating, I want to see it in my own words.")
    bullet(doc, "When I have ten quiet minutes, I want one well-made piece of wisdom content + one good question.")
    bullet(doc, "When I'm a quiet teacher with something to share, I want to publish without selling myself.")
    bullet(doc, "When I'm an admin trying to keep the platform clean, I want to see and edit everything in one place.")

    out = OUT / "02_Target_Audience.docx"; doc.save(str(out)); print(f"wrote {out}")


# ============================================================
# 3. SOLUTION DESIGN
# ============================================================

def build_solution_design():
    doc = Document(); _setup(doc)
    cover(doc, "Solution Design", "Deliverable 03 — Solution Design")

    h2(doc, "Design principles")
    bullet(doc, "Calm by default — no notifications nags, no streak shaming, no engagement loops.")
    bullet(doc, "Listen first — Companion replies in 2-5 sentences, never lectures.")
    bullet(doc, "Roles add power, tiers add depth — superadmin > moderator > creator > customer; standard or premium tier.")
    bullet(doc, "Graceful degradation — every AI surface falls back to a deterministic responder if OPENAI_API_KEY is missing.")
    bullet(doc, "One repo, one deploy — Next.js (web) + FastAPI (api) share one Git repo and one Postgres database.")
    bullet(doc, "Operational trust — moderation, subscriptions, and account lifecycle are first-class admin surfaces, not afterthoughts.")

    h2(doc, "System architecture")
    code(doc,
         "                 ┌─────────────────────────────────────────────────────┐\n"
         "                 │                  Browser / PWA                      │\n"
         "                 │  Next.js 15 App Router · React 19 · Tailwind        │\n"
         "                 │  (+ floating <SupportBot/> on every page)           │\n"
         "                 └──────────────┬──────────────────────────────────────┘\n"
         "                                │  RSC + server actions\n"
         "                 ┌──────────────▼──────────────────────────────────────┐\n"
         "                 │              Next.js server                         │\n"
         "                 │   /api/proxy/[...path]  (forwards to FastAPI with   │\n"
         "                 │   the user's JWT cookie attached as Bearer)         │\n"
         "                 │   /api/auth/logout       (clears cookie)            │\n"
         "                 │   /api/auth/session      (sets cookie from token)   │\n"
         "                 │   middleware.ts          (auth + path forwarding)   │\n"
         "                 └──────────────┬──────────────────────────────────────┘\n"
         "                                │  Bearer JWT (HS256)\n"
         "                 ┌──────────────▼──────────────────────────────────────┐\n"
         "                 │              FastAPI service                        │\n"
         "                 │  routers: auth · users · dashboard · content · ai · │\n"
         "                 │           reflection · skg · notifications ·        │\n"
         "                 │           applications · admin · support            │\n"
         "                 │  ai/companion.py  (OpenAI gpt-4o-mini + fallback)   │\n"
         "                 │  routers/support.py (Glimmora Guide chatbot)        │\n"
         "                 └──────────────┬──────────────────────────────────────┘\n"
         "                                │\n"
         "             ┌──────────────────┼─────────────────────┐\n"
         "             ▼                  ▼                     ▼\n"
         "      ┌────────────┐    ┌────────────────┐     ┌────────────────┐\n"
         "      │  Postgres  │    │  Static assets │     │   OpenAI API   │\n"
         "      │ (Neon / DO │    │  (logo.png,    │     │  (Companion +  │\n"
         "      │  managed)  │    │  uploads/)     │     │  noticings +   │\n"
         "      └────────────┘    └────────────────┘     │  support bot)  │\n"
         "                                               └────────────────┘")

    h2(doc, "Domain model (entities)")
    table(doc,
          header=["Entity", "Purpose", "Key fields"],
          rows=[
              ["User", "Identity + role + profile", "id, username, email, role (customer/creator/moderator/superadmin), is_active"],
              ["Subscription", "Admin-granted tier validity", "user_id, tier (standard/premium), start_date, end_date, note, created_by"],
              ["Series", "OTT title", "id, slug, title, category, tier, published, creator_id, tags"],
              ["Episode", "OTT episode", "id, series_id, slug, video_url, duration_seconds, tier, reflection_prompt"],
              ["WatchProgress", "Resume + continue-watching", "user_id, episode_id, position_seconds, completed"],
              ["Conversation", "Companion thread", "id, user_id, title"],
              ["Message", "Single companion turn", "id, conversation_id, role, content, emotion"],
              ["Reflection", "Journal entry", "id, user_id, content, mood, intensity, tags, episode_id, insights"],
              ["CreatorApplication", "Public apply submission", "id, user_id, full_name, email, pitch, links[], attachments[], status, decided_by, decision_note"],
              ["Notification", "In-app bell feed", "id, user_id, kind, title, body, link, read_at"],
          ])

    h2(doc, "Roles & permissions matrix")
    table(doc,
          header=["Capability", "Customer", "Creator", "Moderator", "Superadmin"],
          rows=[
              ["Sign up / sign in / edit own profile", "✓", "✓", "✓", "✓"],
              ["Use Companion / Journal / Stories", "✓", "✓", "✗ (gated)", "✗ (gated)"],
              ["Apply to become a creator", "✓", "n/a", "n/a", "n/a"],
              ["Publish series + episodes (roadmap — on progression branch)", "✗", "✓", "—", "✓"],
              ["Approve / reject creator applications", "✗", "✗", "✓", "✓"],
              ["See all users + edit / delete any (non-superadmin)", "✗", "✗", "✗", "✓"],
              ["Grant / revoke subscriptions", "✗", "✗", "✗", "✓"],
              ["Create moderators", "✗", "✗", "✗", "✓"],
              ["Default route after login", "/dashboard", "/dashboard", "/moderate/applications", "/admin/customers"],
          ])

    h2(doc, "Tier model")
    body(doc, "Tier is computed from active Subscription rows, not stored on the user. A user is "
              "'premium' if there's a subscription with tier='premium' whose [start_date, end_date] "
              "interval covers now; otherwise 'standard'. The TierBadge in the sidebar reflects this "
              "live. Premium is granted by a Superadmin via the customer edit page → Subscription "
              "manager → Add subscription popup. Self-serve checkout (Stripe) is a roadmap item.")

    h2(doc, "Key user journeys")

    h3(doc, "A. New customer — first 90 seconds")
    numbered(doc, "Visitor opens landing page → Begin gently → /signup.")
    numbered(doc, "Submits username + password (confirm-password gate) + optional email → JWT cookie set → loginAction routes to /onboarding (or /admin/customers / /moderate/applications by role).")
    numbered(doc, "Walks four cards: Welcome → Intention → Focus → Begin. Skip-for-now is always present.")
    numbered(doc, "On Begin, /v1/users/onboard saves preferences + seeds a welcome conversation.")
    numbered(doc, "Lands on /dashboard. Today's three small steps (Arrive / Notice / Reflect) appear.")

    h3(doc, "B. Daily ritual loop")
    numbered(doc, "Step 1 — Arrive → /companion → send any message → done.")
    numbered(doc, "Step 2 — Notice → /watch → play any episode 60+ seconds → done.")
    numbered(doc, "Step 3 — Reflect → /reflect/new → save any entry → done. Streak +1 when all three done.")

    h3(doc, "C. Public creator application")
    numbered(doc, "Visitor clicks 'Apply as creator' in landing nav or footer (works while logged-out too).")
    numbered(doc, "Fills the /apply form: full name, email, username, password, pitch, links[], attachments[].")
    numbered(doc, "Submit → POST /v1/creator-applications creates the User AND the CreatorApplication atomically + returns a JWT.")
    numbered(doc, "Auto-logged-in and redirected to /under-review. (app)/layout gates the user — they cannot escape until decided.")
    numbered(doc, "Moderator opens /moderate/applications → reads pitch + links + attachments → writes optional note → Approve or Reject.")
    numbered(doc, "Approve flips applicant's role to creator, releases the gate, fires application_approved notification. Reject keeps role customer, fires application_rejected notification.")

    h3(doc, "D. Superadmin lifecycle ops")
    numbered(doc, "/admin/customers (or /creators / /moderators) — role-locked tab. Tap a stat tile to filter (Premium / Standard / Pending apps / etc.). Selected tile gets an amber glow ring.")
    numbered(doc, "Row → ✎ Edit → /admin/customers/[id] → edit identity, role, active flag + Subscription manager.")
    numbered(doc, "Add Subscription popup → tier dropdown + start/end date pickers + note → save. Tier flips immediately + notification fires.")
    numbered(doc, "Row → 🗑 Delete → confirm-type-username modal → cascading delete. Blocked for superadmin targets.")
    numbered(doc, "/admin/moderators → Create moderator popup → seeds an account with role=moderator + moderator_promoted notification.")

    h3(doc, "E. Crisis safety")
    numbered(doc, "User message hits the conservative crisis regex (suicide / self-harm / want to die / etc.) in either Companion OR Glimmora Guide.")
    numbered(doc, "Companion: inline 'Stay with someone tonight' card surfaces with region-aware helplines. The conversation continues normally.")
    numbered(doc, "Glimmora Guide: bot's CRISIS HANDLING prompt block fires — leads with warmth, lists helplines directly in the reply, points to Companion.")

    h2(doc, "Glimmora Guide (in-app help bot)")
    body(doc, "A floating amber bubble (lives in root layout — appears on every page including landing "
              "+ auth + admin). Left-click opens a 360×460 glass chat panel; right-click + drag "
              "repositions and persists to localStorage. Panel smart-edges so it never goes off-screen.")
    body(doc, "Backend (/v1/support/chat) is grounded in a hand-curated system prompt covering every "
              "flow, every role, every tier, every notification trigger. Auth is OPTIONAL — when "
              "present, the user's role + tier are added to the context so answers are tailored. "
              "Uses gpt-4o-mini at temperature 0.3 (consistency over creativity). Graceful keyword "
              "fallback when the OpenAI key is missing.")

    h2(doc, "Cross-cutting concerns")
    h3(doc, "Authentication")
    body(doc, "JWT (HS256, 24h) issued by the API on signup/login. Stored in an httpOnly cookie "
              "named glimmora_session by the Next.js layer. Every Next-side server fetch forwards "
              "it as Authorization: Bearer. loginAction (apps/web/src/lib/auth-actions.ts) routes "
              "the user to their correct landing page in a SINGLE redirect, based on role + "
              "onboarding state + pending creator application — no blank flash.")

    h3(doc, "Authorization")
    body(doc, "Server-side per endpoint. Role hierarchy is additive: superadmin > moderator > "
              "creator > customer. The (app)/layout has an early-return for users with a pending "
              "creator application — they're gated on /under-review until decided, regardless of role.")

    h3(doc, "Visual design")
    body(doc, "Warm-neutral palette (ink + glimmer), no cool colors. The body background is an "
              "inlined SVG of heavily-blurred warm bokeh circles (dark mode darker, light mode "
              "cream); a fractal grain overlay prevents the bg from reading as flat. All admin "
              "surfaces use glassmorphic stat tiles (backdrop-blur 22px, translucent bg, top-edge "
              "highlight, warm glow halo). The app shell is chrome-less on desktop — the sidebar "
              "and topbar utility cluster float directly on the bokeh with no container chrome.")

    out = OUT / "03_Solution_Design.docx"; doc.save(str(out)); print(f"wrote {out}")


# ============================================================
# 4. PROJECT DOCUMENTS
# ============================================================

def build_project_documents():
    doc = Document(); _setup(doc)
    cover(doc, "Project Documents", "Deliverable 04 — Project Documents")

    h2(doc, "Project summary")
    body(doc, f"{PROJECT} is a calm AI-powered space for inner work, combining a guided OTT "
              "library, an AI wisdom companion, a reflection journal with a digital twin, an "
              "in-app help chatbot (Glimmora Guide), and a structured admin layer (Customers / "
              f"Creators / Moderators) into one product. Current shipping version: {VERSION}. "
              f"Owner: {OWNER}.")

    h2(doc, "In-scope (current build on main)")

    h3(doc, "Auth + account lifecycle")
    bullet(doc, "Signup with confirm-password gate, login, logout. JWT in httpOnly cookie. Single-redirect login routing by role + onboarding + pending-app state.")
    bullet(doc, "Show / hide password toggle on every password field (signup, login, /apply, create-moderator).")

    h3(doc, "Customer / member-facing")
    bullet(doc, "4-step onboarding (Welcome / Intention / Focus / Begin) with companion welcome message; skippable.")
    bullet(doc, "Dashboard with 'Today's three small steps' (Arrive / Notice / Reflect) + streak.")
    bullet(doc, "AI Companion (gpt-4o-mini) with emotion classification, suggested reflection prompt, conservative crisis surface with region-aware helplines.")
    bullet(doc, "Stories: series + episode catalog, HLS player, watch-progress save + resume, continue-watching surface.")
    bullet(doc, "Reflection: create journal entry, AI 'noticing' synthesis (✦), digital twin trend chart, milestones.")
    bullet(doc, "Profile: edit identity, role display, tier badge (Standard / Premium).")
    bullet(doc, "Public 'Apply as creator' on landing nav + footer + dashboard CTA. Creates account + pending application atomically.")
    bullet(doc, "Under-review gate for users with a pending application — cannot escape until moderator decides.")

    h3(doc, "Moderator surface")
    bullet(doc, "/moderate/applications — list filterable by pending / approved / rejected.")
    bullet(doc, "Review card with pitch + clickable links + attachments. Approve / Reject with optional note. Notification fires to applicant.")

    h3(doc, "Superadmin surface")
    bullet(doc, "Three role-locked tabs: Customers / Creators / Moderators. Each has its own 4-tile hero strip.")
    bullet(doc, "Stat tiles are CLICKABLE FILTERS — tap to narrow the list, tap again to clear. Selected tile gets an amber glow ring.")
    bullet(doc, "Inline ✎ Edit + 🗑 Delete on every row. Delete opens a confirm-type-username modal and cascades to every owned row (reflections, conversations, watch progress, subscriptions, applications, notifications). Superadmin targets are protected.")
    bullet(doc, "Subscription manager on the edit page — Add / Edit popup with tier dropdown + date pickers. Adding flips the user's tier badge immediately and fires a notification.")
    bullet(doc, "Create moderator popup on /admin/moderators.")

    h3(doc, "In-app help")
    bullet(doc, "Glimmora Guide — floating amber bubble bottom-right, draggable on right-click, openable on left-click. Polls /v1/support/chat backed by gpt-4o-mini with a hand-curated 200-line app-knowledge prompt. Crisis-aware. Tested with 23+ realistic questions before shipping.")

    h3(doc, "Notifications")
    bullet(doc, "Bell in topbar with unread count. Polling every 30s. Triggers: application_submitted (mods), application_approved/rejected (applicant), subscription_changed (customer), moderator_promoted (new mod).")

    h3(doc, "Visual system")
    bullet(doc, "Real brand logo (logo.png) in the sidebar + marketing nav + favicon.")
    bullet(doc, "Pictured warm bokeh background + grain overlay + glassmorphic panels + glowing stat tiles + hero glow orbs.")
    bullet(doc, "Chrome-less app shell — sidebar + topbar are floating elements with no container background.")
    bullet(doc, "Dark + light theme with cross-fading bokeh; persistent via localStorage; no flash of wrong theme.")

    h3(doc, "Infrastructure")
    bullet(doc, "Postgres in dev (local) + prod (Neon / DO managed); lifespan idempotently creates tables on first boot.")
    bullet(doc, "DigitalOcean App Platform deployment via .do/app.yaml (auto-deploy on push to main).")
    bullet(doc, "OpenAI API integration verified live — gpt-4o-mini powers Companion + reflection noticings + support bot.")

    h2(doc, "Out of scope (deferred / on progression branch)")
    bullet(doc, "Stripe self-serve subscription checkout — Premium is admin-granted only.")
    bullet(doc, "Real SMTP wiring — settings.smtp_enabled scaffold exists; password-reset / data-export / self-delete features live on the 'progression' branch.")
    bullet(doc, "Email verification at signup.")
    bullet(doc, "Native video upload + transcoding (creators paste URLs today).")
    bullet(doc, "Avatar file uploads (avatarUrl is a text field).")
    bullet(doc, "Native mobile apps (web-first; mobile sidebar + bottom-nav are PWA-ready).")
    bullet(doc, "Automated test suite (manual Playwright walkthroughs done per release).")

    h2(doc, "Milestones")
    table(doc,
          header=["Date", "Tag", "Highlights"],
          rows=[
              ["2026-05-11", "v0.1.0–v0.2.0", "Initial release: auth, content, Companion, journal, circles, onboarding + daily ritual, creator application flow."],
              ["2026-05-12", "v0.2.1–v0.2.2", "Polish: single-redirect signup, role/tier in sidebar, theme on auth pages, motion across buttons/cards/transitions."],
              ["2026-05-15", "v0.3.0", "Feature-completion pass on progression branch: forgot/reset password, change password, data export, self-delete, watch resume + continue-watching, reflection edit + search + range, conversation drawer, creator edit/unpublish/delete, admin search + content moderation + audit log."],
              ["2026-05-16", "v0.4.0-mvp", "Slimmed main to daily-ritual MVP. progression branch preserved with full feature set."],
              ["2026-05-17", "v0.4.5", "Phase A — Spiritual Knowledge Graph seed (12 situations × 5 frameworks × 3 practices, two read endpoints)."],
              ["2026-05-17", "v0.5.0", "Redesign: 4-role hierarchy (customer/creator/moderator/superadmin), public /apply + /under-review gate, subscriptions, notifications, role-locked admin tabs with clickable filter tiles + inline Edit/Delete, Postgres migration, Glimmora Guide help bot, chrome-less floating shell, real brand logo, confirm-password gate."],
          ])

    h2(doc, "Roles & responsibilities")
    table(doc,
          header=["Role", "Owner", "Responsibilities"],
          rows=[
              ["Project Owner / Founder", "Sanjay (Glimmora)", "Vision, prioritisation, brand, creator outreach."],
              ["Engineering", "—", "Frontend (Next.js), backend (FastAPI), DB, deployment."],
              ["Design", "—", "Visual system, motion, copy, calm tone."],
              ["Content / Creator ops", "—", "Curate seed series; review creator applications (Moderator role)."],
              ["QA", "—", "Run TESTING_GUIDE per release; regression Playwright walkthroughs."],
              ["Moderation", "Moderator role-holders", "Review flagged posts (progression branch), decide creator applications."],
          ])

    h2(doc, "Deliverables (this PR)")
    bullet(doc, "All v0.5.0 code on main — typecheck + production build green.")
    bullet(doc, "Tech docs pack (this folder, six documents).")

    h2(doc, "Risks & mitigations")
    table(doc,
          header=["Risk", "Impact", "Mitigation"],
          rows=[
              ["Stripe not yet integrated", "Cannot collect revenue self-serve", "Premium gifted by admin; Stripe is next milestone."],
              ["Single-developer velocity", "Roadmap slippage", "Tight scope per release; reuse progression-branch features when needed."],
              ["AI cost variance per active user", "Margin pressure", "Model is configurable (OPENAI_MODEL); temperature tuned per surface."],
              ["Crisis-detection false negatives", "Safety incident", "Conservative regex by design; explicit CRISIS HANDLING in both Companion + Guide prompts."],
              ["Creator quality drift", "Brand dilution", "Public apply form + hand-vetted by moderators; admin can disable/delete any creator."],
              ["SMTP not wired (no password reset)", "User onboarding friction", "Documented in support bot's responses; superadmin can reset directly in DB if needed. Wire SMTP before public launch."],
          ])

    h2(doc, "Definition of Done (per release)")
    bullet(doc, "TypeScript compiles cleanly (tsc --noEmit) and Next.js production build succeeds.")
    bullet(doc, "Backend imports cleanly and /health returns ok.")
    bullet(doc, "Full role-based Playwright walkthrough passes (customer, creator, moderator, superadmin).")
    bullet(doc, "Glimmora Guide answers 5+ smoke questions accurately.")
    bullet(doc, "USER_MANUAL.md, TESTING_GUIDE.md, and Tech docs reflect the new behavior.")

    out = OUT / "04_Project_Documents.docx"; doc.save(str(out)); print(f"wrote {out}")


# ============================================================
# 5. TEST CASES
# ============================================================

def build_test_cases():
    doc = Document(); _setup(doc)
    cover(doc, "Test Cases", "Deliverable 05 — Test Cases")

    body(doc, f"Test pack for {PROJECT} v{VERSION}. Run against http://127.0.0.1:3000 (web) + "
              "http://127.0.0.1:8000 (api). Bootstrap superadmin: superadmin / 1 (from .env). "
              "Real Postgres in dev; lifespan auto-creates tables on first boot.")

    def case(num, title_, role, pre, steps, expected):
        h3(doc, f"TC-{num:03d}  ·  {title_}")
        body(doc, f"Role: {role}    ·    Pre-conditions: {pre}")
        for s in steps:
            numbered(doc, s)
        p = doc.add_paragraph()
        r = p.add_run("Expected: "); r.bold = True; r.font.size = Pt(11)
        rb = p.add_run(expected); rb.font.size = Pt(11)

    h2(doc, "1.  Auth + account")

    case(1, "Sign up with confirm-password gate",
         "Anonymous", "App reachable.",
         ["Open /signup.",
          "Fill: name=Ren, username=ren_test, email=blank, password=hello1, confirmPassword=hello2.",
          "Observe the 'Create account' button.",
          "Fix confirmPassword to hello1.",
          "Click 'Create account'."],
         "While the two passwords mismatch, the button is disabled and an inline 'Doesn't match…' error shows under the confirm field. When matched, button enables. Submit lands DIRECTLY on /onboarding (single redirect, no flash).")

    case(2, "Password visibility toggle",
         "Anonymous", "On /signup or /login.",
         ["Click the eye icon at the right edge of any password field."],
         "Field text becomes visible; icon flips to EyeOff. Click again to hide. Works on /login, /signup, /apply, and the admin Create-moderator popup.")

    case(3, "Login routes to the right place per role",
         "Anonymous", "Three accounts exist: customer, moderator, superadmin.",
         ["Sign in as superadmin/1.",
          "Sign out, sign in as a moderator.",
          "Sign out, sign in as a customer."],
         "Each lands DIRECTLY on the correct page: superadmin → /admin/customers, moderator → /moderate/applications, customer → /dashboard (or /onboarding if not onboarded). No double-redirect blank flash.")

    h2(doc, "2.  Onboarding")

    case(4, "Four-step onboarding lands on dashboard",
         "Customer (just signed up)", "On /onboarding.",
         ["Step 1 — name + reason → Continue.",
          "Step 2 — intention → Continue.",
          "Step 3 — pick 2 focus areas → Continue.",
          "Step 4 — Begin."],
         "Land on /dashboard. preferences.onboarded=true; intention saved. Companion has 'A first hello' conversation seeded.")

    case(5, "Skip works",
         "Customer (just signed up)", "On /onboarding.",
         ["Click 'Skip for now →' (top right, sits below the floating utility cluster)."],
         "Lands on /dashboard. preferences.onboarded=true with no intention/focus.")

    h2(doc, "3.  Daily ritual + Companion")

    case(6, "Three small steps tick as you complete them",
         "Customer", "Signed in.",
         ["Open /companion → send any message → return to /dashboard.",
          "Open /watch → play any episode 60+s → return to /dashboard.",
          "Open /reflect/new → save any entry → return to /dashboard."],
         "Steps tick in order: Arrive ✓, Notice ✓, Reflect ✓. Streak +1 when all three done.")

    case(7, "Companion replies via real OpenAI",
         "Customer", "OPENAI_API_KEY set.",
         ["Open /companion → type 'I feel restless tonight' → Enter."],
         "Within ~2s an assistant message appears, with emotion chip ('anxious' typically) and a 'question to sit with' card. Falls back to deterministic responder if key missing.")

    case(8, "Crisis safety card surfaces",
         "Customer", "Signed in.",
         ["Type 'I want to die tonight' → Enter."],
         "Rose-bordered safety card appears with India/US-Canada/UK/global helplines. Companion still replies kindly.")

    h2(doc, "4.  Stories")

    case(9, "Watch episode → progress saves → resume",
         "Customer", "Signed in.",
         ["Open any free episode → play 30+ seconds → leave → return to /watch."],
         "Continue-watching row appears at top with the episode + gold progress bar + % watched. Click resumes near saved position.")

    h2(doc, "5.  Reflection + digital twin")

    case(10, "Create reflection with AI noticing",
          "Customer", "OPENAI_API_KEY set.",
          ["/reflect/new → write a sentence → pick mood hopeful → intensity 6 → Save."],
          "Lands on /reflect. Reflections: 1, Day streak: 1, Most-present: hopeful. A ✦ AI-noticing line appears under the entry.")

    h2(doc, "6.  Public creator application + under-review gate")

    case(11, "Anonymous /apply creates account + application",
          "Anonymous", "Logged out.",
          ["Open /apply via header / footer / sidebar 'Apply' link.",
           "Fill name, username, email, password, pitch, optional links.",
           "Submit."],
          "Account created + CreatorApplication added with status=pending. Auto-logged-in. Redirected to /under-review.")

    case(12, "Under-review gate is inescapable",
          "Applicant (customer with pending app)", "On /under-review.",
          ["Try navigating to /dashboard manually.",
           "Try /companion, /reflect, anything in the app."],
          "Server redirects you back to /under-review for every path. No infinite-loop. Layout's pending-application gate early-returns above all other checks.")

    h2(doc, "7.  Moderator")

    case(13, "Moderator sees + approves a pending application",
          "Moderator", "Pending application exists.",
          ["/moderate/applications → filter chip 'pending' → click 'Review and decide' on the row.",
           "Write a short note → Approve."],
          "Tile counters update (Pending −1, Approved +1). Applicant's role flips to creator. Applicant gets application_approved notification.")

    case(14, "Moderator rejects with note",
          "Moderator", "Pending application exists.",
          ["Same flow as TC-13 but click Reject with a note."],
          "Applicant role stays customer. Applicant gets application_rejected notification.")

    h2(doc, "8.  Superadmin")

    case(15, "Three role-locked tabs each show one role only",
          "Superadmin", "Signed in.",
          ["Open /admin/customers — observe rows.",
           "Open /admin/creators — observe rows.",
           "Open /admin/moderators — observe rows."],
          "Each tab shows ONLY the matching role. No filter dropdown — the role IS the page.")

    case(16, "Clickable stat tiles filter the list",
          "Superadmin", "On /admin/customers.",
          ["Click the 'Premium' tile.",
           "Click 'Premium' again to clear.",
           "Click 'Pending apps'."],
          "Selected tile gets an amber glow ring + outline. Table narrows to matching rows. Re-clicking the same tile clears the filter. 'All' is the default state.")

    case(17, "Inline ✎ Edit and 🗑 Delete on every row",
          "Superadmin", "On any of the three admin tabs.",
          ["Click ✎ on a row.",
           "Back to list → click 🗑 on a non-superadmin row.",
           "Type the wrong username in the confirm field → observe button.",
           "Type the right username → Permanently delete."],
          "Edit navigates to /admin/customers/[id]. Delete opens confirm-type-username modal; destructive button is disabled until username matches exactly. On confirm, row vanishes and all owned data is cascaded.")

    case(18, "Delete protections",
          "Superadmin", "Signed in.",
          ["Try DELETE /v1/admin/customers/<your-own-user-id> via API.",
           "Try DELETE on another superadmin row."],
          "First returns 400 'cannot delete self'. Second returns 403 'cannot delete another superadmin'. No row in the UI lets you trigger this from the table — protections are layered.")

    case(19, "Subscription grant flips tier immediately",
          "Superadmin", "Target customer exists.",
          ["/admin/customers → click a customer row.",
           "Subscription manager → Add subscription → tier=premium, today→+30d → save.",
           "As that customer, refresh and look at the sidebar tier badge."],
          "Subscription row appears with 'active' chip. Customer's badge flips from Standard (grey) to Premium (gold). subscription_changed notification fires.")

    case(20, "Create moderator + first login routing",
          "Superadmin → new Moderator", "On /admin/moderators.",
          ["Click 'Create moderator' → fill the popup (full name, username, email, password) → Create.",
           "Sign out. Sign in with the new moderator credentials."],
          "Moderator appears in the list with role=moderator + moderator_promoted notification. Login routes the moderator DIRECTLY to /moderate/applications (not /dashboard).")

    h2(doc, "9.  Notifications")

    case(21, "Bell shows unread + dropdown lists notifications",
          "Anyone", "At least one notification fired for this user.",
          ["Look at the bell icon top-right of the app shell.",
           "Click the bell."],
          "Badge with unread count. Dropdown shows last 30 notifications with title + body + relative time. Click a notification's link to navigate; it marks as read.")

    case(22, "Polling picks up new notifications without refresh",
          "Customer", "Signed in on /dashboard.",
          ["From another window, sign in as superadmin and add a subscription to this customer.",
           "Wait ~30 seconds. Look at the bell."],
          "Bell badge increments without a page refresh.")

    h2(doc, "10.  Glimmora Guide (help bot)")

    case(23, "Bubble opens chat on left-click",
          "Anyone", "On any page.",
          ["Look at the bottom-right of the viewport.",
           "Left-click the amber bubble."],
          "A 360×460 glass chat panel opens, smart-edged to whichever side has more room. Empty state shows 4 suggested starter questions.")

    case(24, "Right-click + drag repositions",
          "Anyone", "Bubble visible.",
          ["Right-click and drag the bubble to a new spot.",
           "Release.",
           "Refresh the page."],
          "Bubble follows the cursor while dragging (no browser context menu appears). Position is constrained to the viewport. After refresh, bubble is in the new spot (persisted in localStorage glimmora-supportbot-pos).")

    case(25, "Bot answers accurately",
          "Anyone", "Bot open.",
          ["Ask: 'How do I become a creator?'",
           "Ask: 'What's the difference between Standard and Premium?'",
           "Ask: 'Who can see my reflections?'"],
          "Q1: explains the /apply form + moderator review. Q2: explains Standard=free default, Premium=admin-granted, no self-serve checkout yet. Q3: only you (and admins, if escalated).")

    case(26, "Bot crisis handling",
          "Anyone", "Bot open.",
          ["Ask: 'I'm thinking of hurting myself.'"],
          "Reply leads with one warm sentence ('I'm really glad you told me.'), surfaces helplines DIRECTLY in the reply (India iCall 9152987821 · US/Canada 988 · UK Samaritans 116 123 · findahelpline.com), points to /companion. No follow-up questions.")

    h2(doc, "11.  Theme + brand")

    case(27, "Theme toggle persists",
          "Anyone", "App open.",
          ["Click moon/sun in the floating top-right cluster.",
           "Refresh."],
          "Theme switches immediately; bokeh background cross-fades between dark and cream variants. Choice persists. No flash of wrong theme.")

    case(28, "Logo renders on every surface",
          "Anyone", "Pages: landing, /signup, /login, /apply, sidebar (any in-app), tab favicon.",
          ["Visually inspect each surface."],
          "logo.png renders cleanly. Tab favicon matches. Brand component scales (sm/md/lg) appropriately by context. Drop-shadow + hover scale visible.")

    out = OUT / "05_Test_Cases.docx"; doc.save(str(out)); print(f"wrote {out}")


# ============================================================
# 6. TECHNICAL DOCUMENTATION
# ============================================================

def build_technical_documentation():
    doc = Document(); _setup(doc)
    cover(doc, "Technical Documentation", "Deliverable 06 — Technical Documentation")

    h2(doc, "Stack")
    bullet(doc, "Frontend: Next.js 15 (App Router, RSC) · React 19 · TypeScript · Tailwind CSS · shadcn-style primitives · jose (JWT verification).")
    bullet(doc, "Backend: FastAPI · SQLAlchemy 2.0 (async) · Alembic · Pydantic v2 · PyJWT · bcrypt · asyncpg.")
    bullet(doc, "Database: Postgres in dev + prod (Neon / DigitalOcean managed). SQLite still supported as a dev fallback via DATABASE_URL.")
    bullet(doc, "AI: OpenAI-compatible (default gpt-4o-mini). Powers Companion + reflection noticings + Glimmora Guide.")
    bullet(doc, "Build / package: pnpm workspaces · turbo · venv for backend.")
    bullet(doc, "Deployment target: DigitalOcean App Platform (.do/app.yaml; auto-deploy on push to main).")

    h2(doc, "Repository layout")
    code(doc,
         "glimmora-one/\n"
         "├── apps/web/                    # Next.js frontend\n"
         "│   ├── public/                  # static assets (logo.png + favicon)\n"
         "│   ├── src/app/                 # routes (App Router)\n"
         "│   │   ├── (app)/               # authed shell with role-aware layout\n"
         "│   │   │   ├── admin/           # superadmin: customers / creators / moderators\n"
         "│   │   │   ├── moderate/        # moderator: applications\n"
         "│   │   │   ├── companion/ watch/ reflect/ profile/ onboarding/\n"
         "│   │   │   └── under-review/    # gate for pending applicants\n"
         "│   │   ├── (auth)/              # public auth pages: login, signup, apply\n"
         "│   │   ├── api/proxy/[...path]/ # forwards to FastAPI w/ Bearer\n"
         "│   │   ├── api/auth/logout/     # clears session cookie\n"
         "│   │   └── api/auth/session/    # sets cookie from token (used by /apply)\n"
         "│   ├── src/components/          # presentational + client components\n"
         "│   │   ├── admin/               # role-locked admin tables + clients\n"
         "│   │   ├── support-bot.tsx      # Glimmora Guide floating widget\n"
         "│   │   ├── notification-bell.tsx\n"
         "│   │   ├── app-shell.tsx        # role-aware sidebar + utility cluster\n"
         "│   │   ├── brand.tsx            # logo (sm/md/lg sizes)\n"
         "│   │   └── ui/                  # primitives (button, input, password-input, ...)\n"
         "│   └── src/lib/                 # backend client, session, types, server actions\n"
         "├── backend/                     # FastAPI service\n"
         "│   ├── app/\n"
         "│   │   ├── routers/             # one per domain\n"
         "│   │   │   ├── auth.py / users.py / dashboard.py / content.py\n"
         "│   │   │   ├── ai.py / reflection.py / skg.py\n"
         "│   │   │   ├── notifications.py / applications.py / admin.py\n"
         "│   │   │   └── support.py       # Glimmora Guide chatbot\n"
         "│   │   ├── ai/companion.py      # Companion AI orchestration\n"
         "│   │   ├── models.py            # SQLAlchemy: User, Subscription, Notification, ...\n"
         "│   │   ├── schemas.py           # Pydantic v2 with camelCase alias generator\n"
         "│   │   ├── services.py          # active_tier(), hydrate_user(), ...\n"
         "│   │   ├── security.py          # bcrypt + JWT\n"
         "│   │   ├── deps.py              # CurrentUser dependency\n"
         "│   │   ├── db.py                # async engine + session factory\n"
         "│   │   ├── bootstrap.py         # superadmin + demo seeds\n"
         "│   │   └── main.py              # app factory + lifespan + create_all\n"
         "│   ├── data/skg.json            # Spiritual Knowledge Graph seed\n"
         "│   └── alembic/                 # migrations (used in prod)\n"
         "├── docs/                        # USER_MANUAL, TESTING_GUIDE, PRODUCT_FLOWS,\n"
         "│   ├── Tech docs/               # this deliverables pack\n"
         "│   └── brand/                   # logo SVGs + preview\n"
         "└── .do/app.yaml                 # DigitalOcean deploy spec")

    h2(doc, "Environment variables")
    table(doc,
          header=["Var", "Purpose", "Default / Example"],
          rows=[
              ["JWT_SECRET", "Shared HS256 signing secret (api + web).", "32+ random chars"],
              ["JWT_EXPIRES_HOURS", "Session lifetime.", "24"],
              ["DATABASE_URL", "SQLAlchemy async URL.", "postgresql+asyncpg://user:pass@host:5432/one"],
              ["BACKEND_URL", "Where Next proxies API calls.", "http://localhost:8000"],
              ["OPENAI_API_KEY", "Enables LLM Companion + AI noticing + Guide.", "(unset → fallback)"],
              ["OPENAI_MODEL", "Override LLM model.", "gpt-4o-mini"],
              ["BOOTSTRAP_SUPERADMIN_USERNAME/_PASSWORD/_EMAIL", "Seed superadmin on first boot.", "superadmin / 1 / admin@glimmora.ai"],
              ["ALLOWED_ORIGINS", "CORS allowlist (CSV).", "http://localhost:3000"],
          ])

    h2(doc, "REST API reference (selected — current main)")
    body(doc, "All endpoints return the envelope {success, data, error}. Auth: Authorization: "
              "Bearer <jwt> (Next.js proxy attaches it from the httpOnly cookie). camelCase on the "
              "wire (Pydantic alias generator).")

    h3(doc, "Auth")
    table(doc, header=["Method + Path", "Purpose"], rows=[
        ["POST /v1/auth/signup", "Create customer account; returns JWT."],
        ["POST /v1/auth/login", "Login by username or email; returns JWT."],
        ["POST /v1/auth/logout", "Stateless; cookie cleared by Next.js layer."],
        ["GET  /v1/auth/me", "Current user with derived role + tier + hasPendingApplication."],
    ])

    h3(doc, "Users + dashboard + content + AI + reflection")
    table(doc, header=["Method + Path", "Purpose"], rows=[
        ["PATCH /v1/users/me", "Update own profile."],
        ["POST  /v1/users/onboard", "Persist 4-step onboarding + seed welcome conversation."],
        ["GET   /v1/dashboard/today", "Three-step state for today + streak."],
        ["GET   /v1/content/series · /series/{slug} · /episodes/{id}", "Public catalog (published only)."],
        ["GET   /v1/content/progress/{episode_id} · POST /v1/content/progress · GET /v1/content/continue-watching", "Watch progress + resume."],
        ["POST  /v1/ai/chat", "Companion turn — returns reply + emotion + suggested reflection + crisis flag."],
        ["GET   /v1/reflection · POST /v1/reflection", "List + create journal entries (with AI noticing)."],
        ["GET   /v1/reflection/insights/twin", "Digital twin snapshot (trend, streak, mood, milestones)."],
        ["GET   /v1/skg/situations · /v1/skg/situations/{slug}", "Read-only SKG (12 situations × 5 frameworks × 3 practices)."],
    ])

    h3(doc, "Notifications, applications, support bot")
    table(doc, header=["Method + Path", "Purpose"], rows=[
        ["GET    /v1/notifications", "List up to 30 most-recent notifications."],
        ["GET    /v1/notifications/unread-count", "Unread count (polled by bell)."],
        ["POST   /v1/notifications/{id}/read · POST /v1/notifications/read-all", "Mark read."],
        ["POST   /v1/creator-applications", "PUBLIC. Creates account + pending application atomically. Returns JWT."],
        ["GET    /v1/creator-applications/me", "Current user's latest application (if any)."],
        ["GET    /v1/moderate/applications?status_filter=...", "Moderator: list (pending/approved/rejected)."],
        ["GET    /v1/moderate/applications/{id}", "Moderator: detail."],
        ["POST   /v1/moderate/applications/{id}/decide", "Moderator: approve or reject + optional note. Fires notification."],
        ["POST   /v1/support/chat", "Glimmora Guide. Auth optional (role-tailored if present)."],
    ])

    h3(doc, "Admin (superadmin only)")
    table(doc, header=["Method + Path", "Purpose"], rows=[
        ["GET    /v1/admin/customers?role=...&q=...", "List users filtered by role + free-text."],
        ["GET    /v1/admin/customers/{id}", "Detail (hydrated with active tier + hasPendingApplication)."],
        ["PATCH  /v1/admin/customers/{id}", "Edit identity, role, active flag."],
        ["DELETE /v1/admin/customers/{id}", "Cascade delete (blocked for superadmin + self)."],
        ["GET    /v1/admin/customers/{id}/subscriptions · POST /v1/admin/customers/{id}/subscriptions", "List + create subscriptions for a user. Notifies the customer."],
        ["PATCH  /v1/admin/subscriptions/{id} · DELETE /v1/admin/subscriptions/{id}", "Edit / delete a subscription."],
        ["GET    /v1/admin/moderators · POST /v1/admin/moderators · DELETE /v1/admin/moderators/{id}", "List / create / demote moderators."],
    ])

    h2(doc, "Authentication + routing")
    bullet(doc, "JWT (HS256) issued by api/security.py:create_access_token. Subject = user.id; extra claim role; iat/exp set.")
    bullet(doc, "Next.js stores it in an httpOnly cookie glimmora_session (sameSite=lax, secure in prod). Server fetches read it via cookies() and forward as Bearer.")
    bullet(doc, "loginAction (apps/web/src/lib/auth-actions.ts) computes the correct landing page ONCE based on role + onboarding state + pending-application — single redirect, no blank flash.")
    bullet(doc, "(app)/layout is the unified routing gate: pending application → /under-review (early-return); else role-based home (superadmin → /admin/customers, moderator → /moderate/applications); else onboarding check for customers/creators.")

    h2(doc, "AI orchestration")
    bullet(doc, "ai/companion.py:respond() — Companion entrypoint. Returns CompanionResult(reply, emotion, reflection_prompt, crisis).")
    bullet(doc, "synthesize_reflection_insight() — single ✦ noticing for each saved reflection.")
    bullet(doc, "routers/support.py — Glimmora Guide. Custom system prompt is the bot's knowledge base (every flow, role, tier, notification trigger). Includes a hard-coded CRISIS HANDLING block that lists region-aware helplines verbatim on self-harm prompts.")
    bullet(doc, "Crisis regex (both Companion + Guide) is conservative — explicit phrases only.")

    h2(doc, "Frontend conventions")
    bullet(doc, "All API calls go through src/lib/backend.ts → fetch BACKEND_URL with the user's Bearer attached on the server, or /api/proxy on the client.")
    bullet(doc, "camelCase boundary: backend serialises snake_case → camelCase via pydantic.alias_generators.to_camel. The UI consumes camelCase only.")
    bullet(doc, "Pages do server fetches; client components handle interactivity (forms, dropdowns, popups, drag).")
    bullet(doc, "AppShell is role-aware — different sidebar items per role; same shell renders nothing on the /onboarding and /under-review focused gates only when those pages opt out.")
    bullet(doc, "SupportBot is mounted in the ROOT layout so it's available on every page, including unauthed surfaces (landing, login, signup, apply).")
    bullet(doc, "Animation: Tailwind keyframes (fade-up, scale-in, gradient-pan, breathe). All animations honour prefers-reduced-motion via globals.css guard.")

    h2(doc, "Deployment")
    body(doc, "Push to main. DigitalOcean reads .do/app.yaml: a PRE_DEPLOY migrate job runs alembic "
              "upgrade head, then the api service (routes /v1/* and /uploads/*) and web service "
              "(routes /*) deploy. One domain, no CORS in prod.")

    h2(doc, "Local development")
    code(doc,
         "# 1. Backend\n"
         "cd backend && python -m venv .venv\n"
         "./.venv/Scripts/pip install -e .            # Windows\n"
         "cp ../.env.example .env                     # then edit\n"
         "./.venv/Scripts/python -m uvicorn app.main:app --port 8000 --host 127.0.0.1\n"
         "\n"
         "# 2. Frontend (in another terminal)\n"
         "pnpm install\n"
         "cp .env.example apps/web/.env.local         # set BACKEND_URL + JWT_SECRET\n"
         "pnpm --filter web dev\n"
         "\n"
         "# OR: from the repo root\n"
         "npm run dev                                  # runs both halves concurrently\n"
         "\n"
         "# Default credentials (created on first boot):\n"
         "#   superadmin / 1")

    h2(doc, "Quality gates per release")
    bullet(doc, "pnpm --filter web typecheck → clean.")
    bullet(doc, "pnpm --filter web build → clean (no missing routes).")
    bullet(doc, "Backend lifespan starts (auto-creates Postgres tables on first boot).")
    bullet(doc, "Full Playwright walkthrough across roles passes.")
    bullet(doc, "Glimmora Guide answers 5+ smoke questions accurately.")
    bullet(doc, "USER_MANUAL.md, TESTING_GUIDE.md, and Tech docs reflect new behavior.")

    out = OUT / "06_Technical_Documentation.docx"; doc.save(str(out)); print(f"wrote {out}")


# ------------------------------------------------------------

if __name__ == "__main__":
    build_usp()
    build_target_audience()
    build_solution_design()
    build_project_documents()
    build_test_cases()
    build_technical_documentation()
