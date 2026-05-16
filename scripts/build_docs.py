"""Build kid-friendly Word docs (TESTING_GUIDE.docx + USER_MANUAL.docx).

Scope: the slim MVP on `main` (v0.4.0-mvp). The full feature-set docs live
on the `progression` branch.

Run from the repo root:
    python scripts/build_docs.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SHOTS = ROOT / "docs" / "screenshots"
OUT = ROOT / "docs"


# ---------------------- helpers ----------------------

def _setup(doc):
    for s in doc.sections:
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x33, 0x2B, 0x2B)
    p.paragraph_format.space_after = Pt(6)


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x44, 0x3C, 0x3C)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0xAD, 0x6B, 0x2D)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(11)


def test(doc, n, title_, you_do, you_see):
    """A numbered test the kid can run."""
    p = doc.add_paragraph()
    r = p.add_run(f"Test {n}. {title_}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x5B)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("👉 You do: ")
    r1.bold = True
    r1.font.size = Pt(11)
    p1.add_run(you_do).font.size = Pt(11)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("👀 You should see: ")
    r2.bold = True
    r2.font.size = Pt(11)
    p2.add_run(you_see).font.size = Pt(11)

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Result: ___ Pass     ___ Fail     Notes: ____________________")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def step(doc, n, title_, do, see):
    """A numbered user-manual step. `do` = what to do, `see` = what to expect."""
    p = doc.add_paragraph()
    r = p.add_run(f"Step {n}. {title_}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x5B)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)

    p1 = doc.add_paragraph()
    r1 = p1.add_run("👉 You do: ")
    r1.bold = True
    r1.font.size = Pt(11)
    p1.add_run(do).font.size = Pt(11)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("👀 You should see: ")
    r2.bold = True
    r2.font.size = Pt(11)
    p2.add_run(see).font.size = Pt(11)


def shot(doc, filename, caption=None, width_in=5.8):
    path = SHOTS / filename
    if not path.exists():
        body(doc, f"[screenshot missing: {filename}]")
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(10)
        cr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def page_break(doc):
    doc.add_page_break()


def callout(doc, title_, body_text):
    p = doc.add_paragraph()
    r = p.add_run("ℹ️  " + title_)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xAD, 0x6B, 0x2D)
    p.paragraph_format.space_before = Pt(8)
    pb = doc.add_paragraph()
    rb = pb.add_run(body_text)
    rb.font.size = Pt(11)


# ============================================================
# USER MANUAL — for the kid using the app
# ============================================================

def build_user_manual():
    doc = Document(); _setup(doc)

    h1(doc, "Glimmora ONE — How to use the app")
    body(doc, "A friendly guide for anyone who can read. If you're 12 and you've never seen this app before, this guide is for you.")
    body(doc, "App version: 0.4.0-mvp   ·   Last updated: 2026-05-16")

    page_break(doc)

    # ---- What it is ----
    h2(doc, "1. What this app is, in one breath")
    body(doc, "Glimmora ONE is a calm room on your phone or computer. There are four small things inside:")
    bullet(doc, "A Companion — a kind AI you can talk to. It listens.")
    bullet(doc, "Stories — short, beautiful videos that make your brain feel quieter.")
    bullet(doc, "A journal — a little place to write a sentence about how you feel.")
    bullet(doc, "A daily ritual — three tiny things to do each day.")
    body(doc, "That's it. No likes. No followers. No notifications buzzing at you. Nobody can see your stuff but you.")

    # ---- Sign up ----
    h2(doc, "2. Making your account")
    body(doc, "This is the first thing you'll do.")
    step(doc, 1, "Open the app",
         "Go to http://127.0.0.1:3000 in a web browser (the page your grown-up tells you to use).",
         "A calm cream / dark page with the words 'A calm room for your inner life.' in the middle.")
    shot(doc, "01-landing.png", "The landing page.")
    step(doc, 2, "Click 'Begin gently'",
         "There's a big button that says 'Begin gently'. Click it.",
         "A small form with four boxes: name, username, email (optional), password.")
    shot(doc, "02-signup.png", "The sign-up page.")
    step(doc, 3, "Fill in the form",
         "Name: anything you like ('Ren'). Username: a short word just you remember. Email: leave it blank. Password: anything ('hello' is fine while testing).",
         "After clicking 'Create account' you land directly on the welcome page. The sidebar on the left shows your name and the word 'member'.")
    callout(doc, "Tip",
            "Use a real password when this is your real account. While testing, even one letter works.")

    # ---- Onboarding ----
    h2(doc, "3. The four hello questions")
    body(doc, "The first time you sign up, the app asks four small questions. Each one has a 'Skip for now →' link in the top-right — you can skip any of them.")
    bullet(doc, "Step 1 — What should we call you? (Just type a name.)")
    bullet(doc, "Step 2 — In one sentence, what do you hope shifts even a little? (Like: 'I'd like to feel less rushed in the mornings.')")
    bullet(doc, "Step 3 — Pick up to four areas that feel alive for you. (Stillness, Becoming, Feeling, Grief, Joy, Relationships, Work, Sleep, Creativity.)")
    bullet(doc, "Step 4 — Click 'Begin'.")
    shot(doc, "03-onboarding-welcome.png", "The first welcome step.")
    body(doc, "When you finish (or skip), the app drops you on the Dashboard. The Companion has already written you a tiny 'A first hello' message you can read in the Companion later.")

    # ---- Dashboard ----
    h2(doc, "4. Today's three small steps")
    body(doc, "Every day when you open the app, the Dashboard shows three tiny things you can do. They are never required.")
    bullet(doc, "1. Arrive — say hi to the Companion (any message at all).")
    bullet(doc, "2. Notice — watch one short Story.")
    bullet(doc, "3. Reflect — write one sentence in your journal.")
    body(doc, "Doing all three turns the streak number up by one. Missing a day is fine. Really. Nobody is keeping score.")
    shot(doc, "04-dashboard-member.png", "The dashboard for a brand-new account.")

    # ---- Companion ----
    h2(doc, "5. Talking to the Companion")
    step(doc, 4, "Open Companion",
         "Click 'Companion' on the left sidebar.",
         "An empty chat with the words 'I'm here.' in the middle and four soft starter sentences as buttons.")
    shot(doc, "06-companion-empty.png", "The empty Companion view.")
    step(doc, 5, "Say something",
         "Click a starter sentence, OR type your own in the box at the bottom and press Enter.",
         "Your message appears in a peach bubble on the right. After a moment, the Companion writes back on the left. Sometimes a small box appears with 'A question to sit with'.")
    callout(doc, "The Companion is NOT a real person",
            "It is an AI built to listen kindly. It is not a therapist and not a doctor. If something serious is happening, see Section 12 — The safety rule.")
    step(doc, 6, "The safety net",
         "If you ever write something serious like 'I want to die', a rose-pink card appears with phone numbers for real people who can help.",
         "A rose-colored card with helplines for India, US/Canada, UK, and a global directory. The Companion still writes back kindly.")
    shot(doc, "07-companion-with-crisis.png", "The safety card.")

    # ---- Stories ----
    h2(doc, "6. Watching a Story")
    step(doc, 7, "Open Stories",
         "Click 'Stories' on the left sidebar.",
         "A library page with three groups of short series.")
    shot(doc, "10-stories-library.png", "The library page.")
    step(doc, 8, "Pick a series, pick an episode",
         "Click any cover. Then click any episode in the list.",
         "The episode page opens with a video player and a quiet 'A question for after' box below it.")
    shot(doc, "11-series-page.png", "Inside a series.")
    step(doc, 9, "Watch — and come back later",
         "Play the video. Pause it whenever you want. Leave the page. Come back another day and click the same episode.",
         "The video starts playing NEAR WHERE YOU STOPPED — not from the beginning. The app remembers for you.")
    step(doc, 10, "Continue watching",
         "When you go back to 'Stories' after watching part of something, look at the very top of the page.",
         "A row called 'Continue watching' appears with the episode you started, a little gold bar showing how far you got, and 'resume' text. Click it to keep going.")

    # ---- Reflect ----
    h2(doc, "7. Writing in your journal (Reflect)")
    step(doc, 11, "Open Reflect",
         "Click 'Reflect' on the left sidebar.",
         "A page called 'Your inner weather.' with four stat cards and a 30-day chart.")
    step(doc, 12, "Write a reflection",
         "Click 'New reflection' in the top right. Type ONE sentence. Tap a mood (like 'hopeful'). Drag the intensity slider. Click 'Save reflection'.",
         "The page jumps back to the journal. The number 'Reflections' goes up by one. Your entry appears at the top. If the AI is on, a small ✦ noticing line appears under it.")
    shot(doc, "08-reflect-new.png", "Writing a reflection.")
    shot(doc, "09-reflect-digest.png", "After saving — stats updated.")

    # ---- Map ----
    h2(doc, "8. Your inner-weather map")
    body(doc, "Below your journal on the same Reflect page, you'll see:")
    bullet(doc, "Reflections — how many you've written, ever.")
    bullet(doc, "Day streak — days in a row you wrote at least one.")
    bullet(doc, "Most-present — the mood you've felt most often.")
    bullet(doc, "Avg. intensity — how strong your feelings have been, on a 1–10 scale.")
    body(doc, "There's a 'Last 30 days' chart. Each bar is one day. Taller bar = bigger feelings. The color tells you what mood. Empty days are faint stubs — that's totally fine.")
    body(doc, "Milestones appear on the right when you reach little markers like 'first reflection logged' or 'a week of noticing'.")

    # ---- Profile ----
    h2(doc, "9. Profile — your tiny page about you")
    body(doc, "Click 'Profile' on the left sidebar. You can change your name, write a short bio, and add an avatar URL. Click 'Save' when you're done.")
    shot(doc, "14-profile-member.png", "The Profile page.")

    # ---- Theme ----
    h2(doc, "10. Dark mode and light mode")
    body(doc, "Look in the top-right of any page. There's a small moon (or sun) button. Click it. The whole app switches between calm dark and soft light. The app remembers your choice for next time.")

    # ---- Help ----
    h2(doc, "11. If something looks weird")
    bullet(doc, "White screen that won't load → refresh the page (F5).")
    bullet(doc, "'Companion is resting' → wait 5 seconds, send the message again.")
    bullet(doc, "Login fails → check the username and password. They care about CAPITAL letters.")
    bullet(doc, "Video shows a black box → try another episode.")
    bullet(doc, "Stuck on the welcome questions → click 'Skip for now →' in the top right.")
    body(doc, "If nothing works, tell your grown-up — they can restart the app for you.")

    # ---- Safety ----
    h2(doc, "12. The safety rule")
    body(doc, "Glimmora is not a doctor or a therapist. It is a calm space to think and feel.")
    body(doc, "If you are in real trouble — if you want to hurt yourself, or someone is hurting you — please tell a real person you trust. A parent. A teacher. A school counselor. Or call one of these numbers any time, day or night:")
    bullet(doc, "India: iCall — 9152987821   ·   Vandrevala Foundation — 1860-2662-345")
    bullet(doc, "US / Canada: 988 (call or text)")
    bullet(doc, "UK / ROI: Samaritans — 116 123")
    bullet(doc, "Anywhere else: findahelpline.com")
    body(doc, "The Companion will show you these numbers too if you write about wanting to hurt yourself. Please use them. You're not alone. Even at 3am.")

    out = OUT / "USER_MANUAL.docx"
    doc.save(str(out))
    print(f"wrote {out}")


# ============================================================
# TESTING GUIDE — for the kid testing the app
# ============================================================

def build_testing_guide():
    doc = Document(); _setup(doc)

    h1(doc, "Glimmora ONE — Testing Guide")
    body(doc, "Pretend you're a detective. Your job is to make sure the app works for everyone. This guide has tests anyone (even a 12-year-old) can run. Each test is a tiny mission — do the thing, check what you see, write ✓ or ✗.")
    body(doc, "App version: 0.4.0-mvp   ·   Last updated: 2026-05-16")

    page_break(doc)

    h2(doc, "Before you start")
    body(doc, "Two things need to be running on your computer:")
    bullet(doc, "The backend at http://127.0.0.1:8000 — if you open it in a browser, you should see some words. If you see an error, ask your grown-up to start it.")
    bullet(doc, "The website at http://127.0.0.1:3000 — opens the Glimmora landing page.")
    body(doc, "You will use two accounts:")
    bullet(doc, "A brand-new person (you make this in Test 3) — for testing as a regular member.")
    bullet(doc, "The boss / superadmin — already exists. Username: superadmin   ·   Password: ChangeMe!2026")

    callout(doc, "How to mark a test",
            "If it does what we expected → write ✓ Pass. If it doesn't → write ✗ Fail and a tiny note about what was different. That's it.")

    page_break(doc)

    # ----- A. Landing -----
    h2(doc, "A. Landing page")
    test(doc, 1, "The landing page loads",
         "Open http://127.0.0.1:3000/ in a browser.",
         "A calm page with the words 'A calm room for your inner life.' in the middle. Two buttons: 'Begin gently' and 'See what's inside'. In the top right, a tiny moon (or sun) icon.")
    shot(doc, "01-landing.png", "The landing page.")
    test(doc, 2, "Scroll-down feels gentle",
         "Scroll the page slowly.",
         "More cards quietly fade in as you scroll. Nothing pops or shouts.")

    # ----- B. Sign up & log in -----
    h2(doc, "B. Sign up & log in")
    test(doc, 3, "You can sign up",
         "Click 'Begin gently'. Fill: Name=Ren, Username=ren_test, Email=blank, Password=hello. Click 'Create account'.",
         "The app jumps STRAIGHT to a welcome page (the four questions). It should NOT show a blank screen first. The sidebar shows 'Ren · member'.")
    shot(doc, "02-signup.png", "The sign-up page.")
    test(doc, 4, "Same username is not allowed twice",
         "Open /signup in a new tab. Try to make another account with username=ren_test.",
         "An error message saying the username is already in use.")
    test(doc, 5, "Log out",
         "In the top-right of any page inside the app, click 'Sign out'.",
         "The app sends you to the landing page. If you type /dashboard in the URL, it sends you to the login page instead.")
    test(doc, 6, "Log back in",
         "Click 'Sign in' on the landing page. Type ren_test and hello. Click 'Sign in'.",
         "You land on the dashboard.")
    test(doc, 7, "Wrong password fails",
         "Sign out. Try to log in with username=ren_test, password=wrong.",
         "A red message saying invalid credentials. You stay on the login page.")

    # ----- C. Onboarding -----
    h2(doc, "C. The four hello questions")
    test(doc, 8, "Walk through all four steps",
         "Sign up as a fresh person (kid_test / hello). Walk through: type a name + a one-line reason → Continue. Type a one-sentence intention → Continue. Pick two squares (like Stillness + Sleep) → Continue. On the last page, click 'Begin'.",
         "You land on the dashboard. Your one-sentence intention is quoted on the dashboard.")
    shot(doc, "03-onboarding-welcome.png", "The first welcome step.")
    test(doc, 9, "Skip works",
         "Sign up as another fresh person. On the welcome page, immediately click 'Skip for now →' in the top right.",
         "You land on the dashboard. No questions asked again next time.")
    test(doc, 10, "Cannot revisit onboarding",
         "After Test 8 or 9, try typing /onboarding in the URL bar and press Enter.",
         "The app sends you back to the dashboard. You only see those questions once.")

    # ----- D. Dashboard -----
    h2(doc, "D. Dashboard (three small steps)")
    test(doc, 11, "Dashboard shows three steps",
         "Open /dashboard.",
         "A greeting, then a card with three small steps: Arrive, Notice, Reflect. All three start with no gold check.")
    shot(doc, "04-dashboard-member.png", "The dashboard for a brand-new member.")
    test(doc, 12, "Step 1 ticks after a Companion message",
         "Click 'Open companion' on the Arrive step. Send any message. Then click 'Today' in the sidebar.",
         "Step 1 — Arrive shows a gold check now.")
    test(doc, 13, "Step 3 ticks after a reflection",
         "Click 'Reflect' in the sidebar. 'New reflection'. Type 'Hello, brain.' Click 'Save reflection'. Return to the dashboard.",
         "Step 3 — Reflect shows a gold check.")

    # ----- E. Companion -----
    h2(doc, "E. Companion")
    test(doc, 14, "Companion page loads empty",
         "Click 'Companion' in the sidebar (in a brand-new account that hasn't used it).",
         "The words 'I'm here.' in the middle. Four soft starter sentences as buttons. A text box at the bottom.")
    shot(doc, "06-companion-empty.png", "The empty Companion view.")
    test(doc, 15, "Click a starter sentence",
         "Click any of the four starters.",
         "Your message appears in a peach bubble on the right. Three little dots animate (the Companion is thinking). Within a few seconds, a reply appears in a soft cream bubble on the left.")
    test(doc, 16, "Type your own message",
         "In the text box, type 'I feel a little tired today.' Press Enter.",
         "Your message appears, then a kind reply. Often a small box also appears with 'A question to sit with' and a question.")
    test(doc, 17, "The safety card appears for serious words",
         "Type a message with crisis words, like 'I want to die tonight'. Press Enter. (This is a test — we're checking the safety net works.)",
         "A rose-pink card with phone numbers for India, US/Canada, UK, and a global website. The Companion still replies kindly.")
    shot(doc, "07-companion-with-crisis.png", "The safety card.")
    test(doc, 18, "Safety card does NOT appear for soft sad words",
         "In a new conversation or new account, type 'I feel a bit lonely tonight.' Press Enter.",
         "The Companion replies kindly. NO pink safety card appears. (That's on purpose — the safety net is for very direct words only.)")

    # ----- F. Stories -----
    h2(doc, "F. Stories")
    test(doc, 19, "Library shows series",
         "Click 'Stories' in the sidebar.",
         "A page called 'The library.' with three groups: emotional intelligence, growth, meditation. Each group has at least one card with a cover picture.")
    shot(doc, "10-stories-library.png", "The library page.")
    test(doc, 20, "Open a series",
         "Click any series cover (try 'Still Mind').",
         "The series page with a cover, a description, and a list of episodes with durations.")
    shot(doc, "11-series-page.png", "Inside a series.")
    test(doc, 21, "Play an episode",
         "Click any episode. Click the play button on the video.",
         "The video starts playing. Below the video, a quiet box with 'A question for after' and a question.")
    test(doc, 22, "Pause + leave + come back resumes",
         "Let the video play for 30 seconds. Pause. Click 'Stories' in the sidebar. Then click the same episode again.",
         "The video starts playing NEAR WHERE YOU STOPPED — not from the very beginning.")
    test(doc, 23, "Continue watching row appears",
         "After Test 22, click 'Stories' in the sidebar.",
         "At the very top of the page, a row called 'Continue watching' with your episode. There's a gold bar showing how much you've watched and 'resume' text.")
    test(doc, 24, "Write a reflection from an episode",
         "Open an episode. Below the video, click 'Write a reflection'. Type a sentence. 'Save reflection'.",
         "A green 'Saved.' line with a link to the journal. Click 'Reflect' in the sidebar — the new entry is there.")

    # ----- G. Reflect -----
    h2(doc, "G. Reflect (journal + map)")
    test(doc, 25, "Empty journal looks gentle",
         "Sign up as a brand-new person, skip onboarding, and click 'Reflect'.",
         "Stats all say zero or '—'. The journal area says 'Nothing yet. The first sentence is the hardest.'")
    test(doc, 26, "Write a reflection",
         "Click 'New reflection'. Type one sentence. Tap the mood 'hopeful'. Drag intensity to 6. Click 'Save reflection'.",
         "Page jumps to /reflect. Reflections: 1. Day streak: 1. Most-present: hopeful. Avg. intensity: 6.0. Today's chip in the 'Last 30 days' chart is colored. Your entry shows up in the journal.")
    shot(doc, "08-reflect-new.png", "Writing a reflection.")
    shot(doc, "09-reflect-digest.png", "After saving.")
    test(doc, 27, "Second reflection adds to stats",
         "Write another reflection with mood 'anxious' and intensity 8.",
         "Reflections: 2. Avg. intensity: 7.0. Most-present may stay 'hopeful' or flip to 'anxious'.")
    test(doc, 28, "Milestones show up",
         "After writing your first reflection, look at the 'Milestones' card on the right.",
         "At least the milestone 'First reflection logged' with a glimmer ✦ next to it.")
    test(doc, 29, "Empty days are faint",
         "Look at the 'Last 30 days' chart.",
         "Today (the right-most stub) is colored. The days before (when you didn't write) are very faint. That's correct.")

    # ----- H. Profile -----
    h2(doc, "H. Profile")
    test(doc, 30, "Profile loads",
         "Click 'Profile' in the sidebar.",
         "Your name as a big serif heading. Your email under it. A card called 'You' with editable name and bio.")
    shot(doc, "14-profile-member.png", "The Profile page.")
    test(doc, 31, "Edit name",
         "Change your name to something new (like 'Tiger'). Click 'Save'.",
         "Name updates at the top of the Profile page. The sidebar also updates.")
    test(doc, 32, "Edit bio",
         "Type something into the bio box. Click 'Save'.",
         "A quiet confirmation that it saved (no error message).")

    # ----- I. Theme -----
    h2(doc, "I. Dark / light mode")
    test(doc, 33, "Theme toggle",
         "Click the moon/sun in the top right of any page. Then refresh the page (F5).",
         "Colors flip immediately. After refresh, the new color stays. No flash of the wrong color.")

    # ----- J. Little things -----
    h2(doc, "J. The little things")
    test(doc, 34, "Protected pages need a login",
         "Sign out. Try opening /dashboard directly.",
         "App sends you to /login?next=/dashboard. After signing in, you land on /dashboard.")
    test(doc, 35, "Mobile sidebar at the bottom",
         "Open the app on your phone, or shrink your browser window narrow.",
         "The side sidebar disappears. Five icons appear at the BOTTOM of the screen — that's the mobile menu.")
    test(doc, 36, "Theme on auth pages",
         "Click the moon/sun on /login and on /signup.",
         "The theme toggle works there too.")

    # ----- Done -----
    page_break(doc)
    h2(doc, "Done? Here's the scorecard")
    body(doc, "Count your ✓s and ✗s. There are 36 tests total.")
    bullet(doc, "Section A — Landing (2 tests)")
    bullet(doc, "Section B — Sign up & log in (5 tests)")
    bullet(doc, "Section C — Onboarding (3 tests)")
    bullet(doc, "Section D — Dashboard (3 tests)")
    bullet(doc, "Section E — Companion (5 tests)")
    bullet(doc, "Section F — Stories (6 tests)")
    bullet(doc, "Section G — Reflect (5 tests)")
    bullet(doc, "Section H — Profile (3 tests)")
    bullet(doc, "Section I — Dark / light (1 test)")
    bullet(doc, "Section J — Little things (3 tests)")
    body(doc, "All ✓ → the app is healthy. 🎉")
    body(doc, "Some ✗ → write down which test number(s) and what you saw. Give that list to your grown-up.")

    h2(doc, "Common gotchas")
    bullet(doc, "Blank page after signing up → dev server is caching old code; restart it.")
    bullet(doc, "'JWT_SECRET is not configured' → ask your grown-up to set up .env.local.")
    bullet(doc, "Login works but the dashboard is empty → use http://127.0.0.1:3000 (not localhost, not https).")
    bullet(doc, "Companion never replies → either OPENAI_API_KEY isn't set, or wait — the simple fallback replies should still appear.")
    bullet(doc, "Video shows a black box → the internet is blocking the test stream; try another episode.")
    bullet(doc, "Backend won't start (port 8000 busy) → restart your computer if you can't find the old process.")

    h2(doc, "When you're done testing")
    body(doc, "If everything worked, tell your grown-up: 'All 36 tests passed.'")
    body(doc, "If something didn't work, write down: 1) the test number (like 'Test 22'), 2) what you saw (like 'the video started from zero'), 3) what you expected to see (like 'it should resume').")
    body(doc, "Good detective work! 🕵️")

    out = OUT / "TESTING_GUIDE.docx"
    doc.save(str(out))
    print(f"wrote {out}")


if __name__ == "__main__":
    build_user_manual()
    build_testing_guide()
